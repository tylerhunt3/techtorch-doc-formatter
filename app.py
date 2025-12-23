import streamlit as st
from docx import Document as ReadDocument
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import xml.etree.ElementTree as ET
import io
import re
import zipfile
from datetime import datetime

# ============================================================================
# CONFIGURATION - TechTorch Standards
# ============================================================================

COLORS = {
    "HEADING1": RGBColor(0x1F, 0x4E, 0x79),
    "HEADING2": RGBColor(0x2E, 0x75, 0xB6),
    "HEADING3": RGBColor(0x40, 0x40, 0x40),
    "BODY": RGBColor(0x00, 0x00, 0x00),
    "SECONDARY": RGBColor(0x66, 0x66, 0x66),
    "CODE_TEXT": RGBColor(0x2E, 0x2E, 0x2E),
    "WHITE": RGBColor(0xFF, 0xFF, 0xFF),
    "TABLE_HEADER_BG": "1F4E79",
    "TABLE_BORDER": "CCCCCC",
    "CODE_BG": "F5F5F5",
    "CODE_BORDER": "BFBFBF",
    "CODE_ACCENT": "1F4E79",
}

SIZES = {
    "TITLE": Pt(24),
    "SUBTITLE": Pt(14),
    "HEADING1": Pt(12),
    "HEADING2": Pt(11),
    "HEADING3": Pt(10),
    "BODY": Pt(9),
    "CODE": Pt(8),
    "HEADER_FOOTER": Pt(8),
}


# ============================================================================
# CONTENT EXTRACTION - Enhanced to handle text boxes and code blocks
# ============================================================================

def extract_text_from_xml_element(element):
    """Extract all text from an XML element and its children."""
    texts = []
    # Handle the namespace prefix for Word XML
    for text_elem in element.iter():
        if text_elem.tag.endswith('}t') and text_elem.text:
            texts.append(text_elem.text)
    return ''.join(texts)


def extract_textboxes_from_docx(docx_bytes):
    """Extract text boxes (code blocks) from the document XML."""
    textboxes = []
    
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            if 'word/document.xml' in zf.namelist():
                xml_content = zf.read('word/document.xml')
                
                # Parse XML
                root = ET.fromstring(xml_content)
                
                # Find all text box content - look for txbx elements
                # The namespace for wordprocessingShape
                wps_ns = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}'
                w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                
                for txbx in root.iter(f'{wps_ns}txbx'):
                    lines = []
                    for para in txbx.iter(f'{w_ns}p'):
                        para_texts = []
                        for t in para.iter(f'{w_ns}t'):
                            if t.text:
                                para_texts.append(t.text)
                        if para_texts:
                            lines.append(''.join(para_texts))
                    
                    if lines:
                        full_text = '\n'.join(lines)
                        if is_code_block(full_text):
                            textboxes.append({
                                "type": "code_block",
                                "lines": lines,
                                "text": full_text
                            })
    except Exception as e:
        st.warning(f"Note: Could not extract text boxes: {str(e)}")
    
    return textboxes


def is_code_block(text):
    """Determine if text looks like a code block."""
    code_indicators = [
        'SELECT', 'FROM', 'WHERE', 'INSERT', 'UPDATE', 'DELETE',  # SQL
        'VAR', 'RETURN', 'IF', 'THEN', 'ELSE',  # DAX/Programming
        'def ', 'class ', 'import ', 'function',  # Python/JS
        '= {', '=> ', '};', '();',  # Code syntax
    ]
    
    text_upper = text.upper()
    matches = sum(1 for indicator in code_indicators if indicator.upper() in text_upper)
    
    # Also check for indentation patterns
    lines = text.split('\n')
    indented_lines = sum(1 for line in lines if line.startswith('    ') or line.startswith('\t'))
    
    return matches >= 2 or indented_lines >= 3


def is_section_header(text, is_bold=False, color=None):
    """Determine if text is a section header."""
    text = text.strip()
    
    # Check for numbered section patterns
    level1_pattern = r'^(\d+)\.\s+[A-Z]'
    level2_pattern = r'^(\d+)\.(\d+)\s+[A-Z]'
    level3_pattern = r'^(\d+)\.(\d+)\.(\d+)\s+[A-Z]'
    
    if re.match(level3_pattern, text):
        return "heading3"
    if re.match(level2_pattern, text):
        return "heading2"
    if re.match(level1_pattern, text):
        return "heading1"
    
    # Check for common header words without numbers
    header_keywords = ['Summary', 'Conclusion', 'Overview', 'Introduction', 
                       'Background', 'Results', 'Analysis', 'Recommendations',
                       'Final Summary', 'Next Steps']
    
    for keyword in header_keywords:
        if text == keyword or text.startswith(keyword + ':'):
            return "heading1"
    
    return None


def is_bullet_point(para, text):
    """Determine if paragraph is a bullet point."""
    # Check style name
    if para.style and 'list' in para.style.name.lower():
        return True
    
    # Check for bullet characters at start
    if text.startswith(('•', '-', '*', '–')):
        return True
    
    # Check for "Result:" pattern which should be a bullet
    if text.startswith('Result:'):
        return True
    
    # Check XML for numbering
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            return True
    
    return False


def extract_content_from_docx(uploaded_file):
    """Extract all content from uploaded Word document."""
    # Read file bytes for text box extraction
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # Reset for python-docx
    
    # Extract text boxes (code blocks)
    textboxes = extract_textboxes_from_docx(file_bytes)
    
    # Read document with python-docx
    doc = ReadDocument(uploaded_file)
    content = []
    
    # Track which textbox we're looking for
    textbox_index = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this paragraph mentions a query (insert code block after it)
        mentions_query = 'following' in text.lower() and ('query' in text.lower() or 'soql' in text.lower())
        
        # Get formatting info
        is_bold = para.runs and para.runs[0].bold if para.runs else False
        style_name = para.style.name.lower() if para.style else ""
        
        # Determine content type
        if 'title' in style_name:
            content.append({"type": "title", "text": text})
        elif is_bullet_point(para, text):
            clean_text = re.sub(r'^[•\-*–]\s*', '', text)
            content.append({"type": "bullet", "text": clean_text})
        else:
            header_type = is_section_header(text, is_bold)
            if header_type:
                content.append({"type": header_type, "text": text})
            else:
                content.append({"type": "paragraph", "text": text})
        
        # Insert code block if this paragraph mentions a query
        if mentions_query and textbox_index < len(textboxes):
            content.append(textboxes[textbox_index])
            textbox_index += 1
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not row_data or cell_text != row_data[-1]:
                    row_data.append(cell_text)
            if row_data:
                table_data.append(row_data)
        if table_data:
            content.append({"type": "table", "data": table_data})
    
    return content


# ============================================================================
# DOCUMENT CREATION - Full TechTorch formatting
# ============================================================================

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    
    tcPr.append(tcBorders)


def create_code_block_table(doc, code_lines):
    """Create a formatted code block as a table."""
    table = doc.add_table(rows=len(code_lines), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    for i, line in enumerate(code_lines):
        cell = table.rows[i].cells[0]
        
        # Set cell background
        set_cell_shading(cell, COLORS["CODE_BG"])
        
        # Set borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        # Left border - thick accent
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:color'), COLORS["CODE_ACCENT"])
        tcBorders.append(left)
        
        # Right border
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), '4')
        right.set(qn('w:color'), COLORS["CODE_BORDER"])
        tcBorders.append(right)
        
        # Top border (only first row)
        if i == 0:
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')
            top.set(qn('w:color'), COLORS["CODE_BORDER"])
            tcBorders.append(top)
        
        # Bottom border (only last row)
        if i == len(code_lines) - 1:
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), COLORS["CODE_BORDER"])
            tcBorders.append(bottom)
        
        tcPr.append(tcBorders)
        
        # Add text
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Twips(20)
        para.paragraph_format.space_after = Twips(20)
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = SIZES["CODE"]
        run.font.color.rgb = COLORS["CODE_TEXT"]
    
    return table


def create_data_table(doc, table_data):
    """Create a formatted data table."""
    if not table_data:
        return None
    
    num_rows = len(table_data)
    num_cols = len(table_data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= len(row.cells):
                continue
            cell = row.cells[col_idx]
            cell.text = cell_text
            
            set_cell_borders(cell, COLORS["TABLE_BORDER"])
            
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Aptos'
                    run.font.size = SIZES["BODY"]
            
            if row_idx == 0:
                set_cell_shading(cell, COLORS["TABLE_HEADER_BG"])
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = COLORS["WHITE"]
    
    return table


def create_formatted_document(content, doc_title, organization="TechTorch Inc."):
    """Create a professionally formatted Word document."""
    doc = ReadDocument()
    
    # Configure default styles
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = SIZES["BODY"]
    style.font.color.rgb = COLORS["BODY"]
    
    # Configure heading styles
    for style_name, size, color in [
        ('Heading 1', SIZES["HEADING1"], COLORS["HEADING1"]),
        ('Heading 2', SIZES["HEADING2"], COLORS["HEADING2"]),
        ('Heading 3', SIZES["HEADING3"], COLORS["HEADING3"]),
    ]:
        heading_style = doc.styles[style_name]
        heading_style.font.name = 'Aptos'
        heading_style.font.size = size
        heading_style.font.color.rgb = color
        heading_style.font.bold = True
    
    # ---- TITLE PAGE ----
    for _ in range(4):
        doc.add_paragraph()
    
    # Document title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(doc_title)
    title_run.font.name = 'Aptos'
    title_run.font.size = SIZES["TITLE"]
    title_run.font.bold = True
    title_run.font.color.rgb = COLORS["BODY"]
    
    # Find subtitle from content
    subtitle_text = None
    for item in content:
        if item["type"] == "title" and item["text"] != doc_title:
            subtitle_text = item["text"]
            break
        elif item["type"] == "heading1":
            subtitle_text = item["text"]
            break
    
    if subtitle_text and subtitle_text != doc_title:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(subtitle_text)
        subtitle_run.font.name = 'Aptos'
        subtitle_run.font.size = SIZES["SUBTITLE"]
        subtitle_run.font.color.rgb = COLORS["BODY"]
    
    # Organization
    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_para.add_run(organization)
    org_run.font.name = 'Aptos'
    org_run.font.size = Pt(12)
    org_run.font.italic = True
    org_run.font.color.rgb = COLORS["BODY"]
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"As of {datetime.now().strftime('%B %Y')}")
    date_run.font.name = 'Aptos'
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = COLORS["BODY"]
    
    # Version
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run("Version 1.0")
    version_run.font.name = 'Aptos'
    version_run.font.size = SIZES["BODY"]
    version_run.font.color.rgb = COLORS["SECONDARY"]
    
    # Page break
    doc.add_page_break()
    
    # ---- MAIN CONTENT ----
    skip_next_heading = subtitle_text is not None
    
    for item in content:
        item_type = item["type"]
        
        if item_type == "title":
            continue
        
        elif item_type == "heading1":
            if skip_next_heading:
                skip_next_heading = False
                continue
            para = doc.add_paragraph(item["text"], style='Heading 1')
            para.paragraph_format.space_before = Twips(300)
            para.paragraph_format.space_after = Twips(100)
        
        elif item_type == "heading2":
            para = doc.add_paragraph(item["text"], style='Heading 2')
            para.paragraph_format.space_before = Twips(200)
            para.paragraph_format.space_after = Twips(80)
        
        elif item_type == "heading3":
            para = doc.add_paragraph(item["text"], style='Heading 3')
            para.paragraph_format.space_before = Twips(160)
            para.paragraph_format.space_after = Twips(60)
        
        elif item_type == "paragraph":
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Twips(160)
            run = para.add_run(item["text"])
            run.font.name = 'Aptos'
            run.font.size = SIZES["BODY"]
            run.font.color.rgb = COLORS["BODY"]
        
        elif item_type == "bullet":
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Twips(720)
            run = para.add_run(item["text"])
            run.font.name = 'Aptos'
            run.font.size = SIZES["BODY"]
            run.font.color.rgb = COLORS["BODY"]
        
        elif item_type == "code_block":
            doc.add_paragraph()
            create_code_block_table(doc, item["lines"])
            doc.add_paragraph()
        
        elif item_type == "table":
            doc.add_paragraph()
            create_data_table(doc, item["data"])
            doc.add_paragraph()
    
    # ---- END OF DOCUMENT ----
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_para.paragraph_format.space_before = Twips(400)
    end_run = end_para.add_run("End of Document")
    end_run.font.name = 'Aptos'
    end_run.font.size = SIZES["BODY"]
    end_run.font.italic = True
    end_run.font.color.rgb = COLORS["SECONDARY"]
    
    return doc


def save_doc_to_bytes(doc):
    """Save document to bytes for download."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================================
# STREAMLIT APPLICATION
# ============================================================================

st.set_page_config(
    page_title="TechTorch Document Formatter",
    page_icon="📄",
    layout="centered"
)

# Header
st.title("📄 TechTorch Document Formatter")
st.markdown("Upload a Word document to automatically apply TechTorch formatting standards.")
st.markdown("*Supports code blocks, tables, section headers, and bullet points.*")

st.divider()

# File upload
uploaded_file = st.file_uploader(
    "Upload your Word document",
    type=["docx"],
    help="Upload a .docx file to format"
)

# Document title input
doc_title = st.text_input(
    "Document Title",
    placeholder="Enter the document title...",
    help="This will appear on the title page"
)

# Organization input
organization = st.text_input(
    "Organization",
    value="TechTorch Inc.",
    help="Organization name for the title page"
)

st.divider()

# Process button
if uploaded_file is not None and doc_title:
    if st.button("Format Document", type="primary"):
        with st.spinner("Processing your document..."):
            try:
                # Extract content
                content = extract_content_from_docx(uploaded_file)
                
                # Show what was extracted
                code_blocks = sum(1 for c in content if c["type"] == "code_block")
                tables = sum(1 for c in content if c["type"] == "table")
                headings = sum(1 for c in content if c["type"].startswith("heading"))
                bullets = sum(1 for c in content if c["type"] == "bullet")
                
                st.info(f"Extracted: {headings} headings, {bullets} bullet points, {code_blocks} code blocks, {tables} tables")
                
                # Create formatted document
                formatted_doc = create_formatted_document(
                    content, 
                    doc_title, 
                    organization
                )
                
                # Save to bytes
                doc_bytes = save_doc_to_bytes(formatted_doc)
                
                # Success message
                st.success("Document formatted successfully!")
                
                # Download button
                clean_title = re.sub(r'[^\w\s-]', '', doc_title).replace(' ', '_')
                filename = f"{clean_title}_Formatted.docx"
                
                st.download_button(
                    label="Download Formatted Document",
                    data=doc_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"Error processing document: {str(e)}")
                st.exception(e)

elif uploaded_file is None:
    st.info("Please upload a Word document to get started.")
elif not doc_title:
    st.info("Please enter a document title.")

# Footer
st.divider()
st.markdown(
    "<div style='text-align: center; color: #666666; font-size: 12px;'>"
    "TechTorch Documentation Formatting Tool v2.0<br>"
    "Now with code block and table support"
    "</div>",
    unsafe_allow_html=True
)
